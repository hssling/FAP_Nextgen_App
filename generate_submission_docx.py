
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_submission_docx():
    # Read the markdown content
    with open('d:/FAP App/FAP_NextGen/AI_CHALLENGE_SUBMISSION.md', 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Title
    title = doc.add_heading('AI Challenge Competition Submission', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('FAP NextGen - Advanced Family Adoption Programme Management')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_page_break()

    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        if line.startswith('## '):
            # Question Header
            header_text = line.replace('## ', '')
            doc.add_heading(header_text, level=1)
        
        elif line.startswith('### '):
            # Sub Header
            sub_header_text = line.replace('### ', '')
            doc.add_heading(sub_header_text, level=2)
            
        elif line.startswith('**') and line.endswith('**'):
             # Bold line / Question Text
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
            
        elif line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
            
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # Numbered List (Simple heuristic)
            text = line.split('. ', 1)[1] if '. ' in line else line
            p = doc.add_paragraph(text, style='List Number')
            
        else:
            # Normal text
            # Handle inline bolding **text**
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are inside ** **
                    run.bold = True

    output_path = 'd:/FAP App/FAP_NextGen/AI_CHALLENGE_SUBMISSION.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_submission_docx()
