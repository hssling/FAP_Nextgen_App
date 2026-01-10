
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_kannada_documentation():
    doc = Document()

    # Title Page
    title = doc.add_heading('FAP NextGen ಆಪ್', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('ಸಮಗ್ರ ದಾಖಲಾತಿ ಮತ್ತು ಪಾತ್ರ-ಆಧಾರಿತ (Role-Based) ಬಳಕೆದಾರರ ಕೈಪಿಡಿ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n' * 5)
    
    details = doc.add_paragraph('ಯಾರಿಗಾಗಿ: MBBS ವಿದ್ಯಾರ್ಥಿಗಳು, ಬೋಧಕರು (ಮೆಂಟರ್‌ಗಳು) ಮತ್ತು ಅಡ್ಮಿನಿಸ್ಟ್ರೇಟರ್‌ಗಳು\nಸಂದರ್ಭ: ಸಾಮರ್ಥ್ಯ ಆಧಾರಿತ ವೈದ್ಯಕೀಯ ಶಿಕ್ಷಣ (CBME)\nಕುಟುಂಬ ದತ್ತು ಕಾರ್ಯಕ್ರಮ (FAP)')
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('ವಿಷಯ ಸೂಚಿ', level=1)
    doc.add_paragraph('1. ಪರಿಚಯ ಮತ್ತು ಉದ್ದೇಶಗಳು')
    doc.add_paragraph('2. ಪಾತ್ರಗಳು ಮತ್ತು ಜವಾಬ್ದಾರಿಗಳ ಅವಲೋಕನ')
    doc.add_paragraph('3. ವಿದ್ಯಾರ್ಥಿ ಕೆಲಸದ ಹರಿವು (ಮೌಲ್ಯಮಾಪನಗಳು ಮತ್ತು ಚಟುವಟಿಕೆಗಳು)')
    doc.add_paragraph('4. ಮೆಂಟರ್/ಶಿಕ್ಷಕ ಕೆಲಸದ ಹರಿವು')
    doc.add_paragraph('5. ಅಡ್ಮಿನಿಸ್ಟ್ರೇಟರ್ ಕೆಲಸದ ಹರಿವು')
    doc.add_paragraph('6. ವೈಶಿಷ್ಟ್ಯಗಳ ಪಟ್ಟಿ')
    
    doc.add_page_break()

    # Section 1: Introduction
    doc.add_heading('1. ಪರಿಚಯ ಮತ್ತು ಉದ್ದೇಶಗಳು', level=1)
    p = doc.add_paragraph(' ')
    p.add_run('FAP NextGen ಆಪ್').bold = True
    p.add_run(' ಎನ್ನುವುದು ರಾಷ್ಟ್ರೀಯ ವೈದ್ಯಕೀಯ ಆಯೋಗದ (NMC) ಮಾರ್ಗಸೂಚಿಗಳ ಪ್ರಕಾರ ಕುಟುಂಬ ದತ್ತು ಕಾರ್ಯಕ್ರಮವನ್ನು (FAP) ಕಾರ್ಯಗತಗೊಳಿಸಲು ವಿನ್ಯಾಸಗೊಳಿಸಲಾದ ಸಮಗ್ರ ಡಿಜಿಟಲ್ ಪ್ಲಾಟ್‌ಫಾರ್ಮ್ ಆಗಿದೆ. ಇದು ವೈದ್ಯಕೀಯ ವಿದ್ಯಾರ್ಥಿಗಳು, ಸಮುದಾಯ ಮತ್ತು ಬೋಧಕ ಮೆಂಟರ್‌ಗಳ ನಡುವೆ ಸೇತುವೆಯಾಗಿ ಕಾರ್ಯನಿರ್ವಹಿಸುತ್ತದೆ.')
    
    doc.add_heading('ಪ್ರಮುಖ ಉದ್ದೇಶಗಳು:', level=2)
    doc.add_paragraph('1. ವೈದ್ಯಕೀಯ ವಿದ್ಯಾರ್ಥಿಗಳಿಂದ ಗ್ರಾಮೀಣ ಕುಟುಂಬಗಳ ದೀರ್ಘಾವಧಿಯ ಆರೋಗ್ಯ ಮೇಲ್ವಿಚಾರಣೆಯನ್ನು ಸುಗಮಗೊಳಿಸುವುದು.', style='List Bullet')
    doc.add_paragraph('2. ಸಮುದಾಯ ರೋಗನಿರ್ಣಯ ಮತ್ತು ಆರೋಗ್ಯ ಹಸ್ತಕ್ಷೇಪಗಳಿಗಾಗಿ ನೈಜ-ಸಮಯದ (real-time) ಡೇಟಾವನ್ನು ಒದಗಿಸುವುದು.', style='List Bullet')
    doc.add_paragraph('3. ಮೆಂಟರ್‌ಗಳು ವಿದ್ಯಾರ್ಥಿಗಳ ಸಾಮರ್ಥ್ಯಗಳನ್ನು ದೂರದಿಂದಲೇ ಮತ್ತು ಪರಿಣಾಮಕಾರಿಯಾಗಿ ಮೌಲ್ಯಮಾಪನ ಮಾಡಲು ಅನುವು ಮಾಡಿಕೊಡುವುದು.', style='List Bullet')

    # Section 2: Roles
    doc.add_heading('2. ಪಾತ್ರಗಳು ಮತ್ತು ಜವಾಬ್ದಾರಿಗಳ ಅವಲೋಕನ', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ಪಾತ್ರ (Role)'
    hdr_cells[1].text = 'ಪ್ರಾಥಮಿಕ ಜವಾಬ್ದಾರಿ'
    
    roles = [
        ('ವಿದ್ಯಾರ್ಥಿ (Student)', 'ಕುಟುಂಬಗಳನ್ನು ದತ್ತು ಪಡೆಯುವುದು, ಭೇಟಿಗಳನ್ನು ನಡೆಸುವುದು, ಡೇಟಾ ಸಂಗ್ರಹಿಸುವುದು ಮತ್ತು ರಿಫ್ಲೆಕ್ಷನ್ (ಚಿಂತನೆ) ಬರೆಯುವುದು.'),
        ('ಮೆಂಟರ್ (ಶಿಕ್ಷಕರು)', 'ವಿದ್ಯಾರ್ಥಿಗಳಿಗೆ ಮಾರ್ಗದರ್ಶನ ನೀಡುವುದು, ರಿಫ್ಲೆಕ್ಷನ್‌ಗಳನ್ನು ಪರಿಶೀಲಿಸುವುದು ಮತ್ತು ಕಾರ್ಯಕ್ಷಮತೆಯನ್ನು ಮೌಲ್ಯಮಾಪನ ಮಾಡುವುದು.'),
        ('ಅಡ್ಮಿನಿಸ್ಟ್ರೇಟರ್', 'ಬಳಕೆದಾರರನ್ನು (ವಿದ್ಯಾರ್ಥಿಗಳು/ಶಿಕ್ಷಕರು) ನಿರ್ವಹಿಸುವುದು, ಸಿಸ್ಟಮ್ ಆರೋಗ್ಯವನ್ನು ನೋಡಿಕೊಳ್ಳುವುದು.')
    ]
    
    for role, resp in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = resp

    doc.add_page_break()

    # Section 3: Student Workflow
    doc.add_heading('3. ವಿದ್ಯಾರ್ಥಿ ಕೆಲಸದ ಹರಿವು (ಮೌಲ್ಯಮಾಪನಗಳು ಮತ್ತು ಚಟುವಟಿಕೆಗಳು)', level=1)
    doc.add_paragraph('ವಿದ್ಯಾರ್ಥಿಗಳು ಪ್ರಾಥಮಿಕ ಡೇಟಾ ಸಂಗ್ರಹಕಾರರು. ಅವರ ಕೆಲಸ ಈ ಕೆಳಗಿನ ಹಂತಗಳನ್ನು ಒಳಗೊಂಡಿರುತ್ತದೆ:')
    
    doc.add_heading('3.1 ವಿದ್ಯಾರ್ಥಿಗಳು ಏನನ್ನು ಮೌಲ್ಯಮಾಪನ (Assess) ಮಾಡುತ್ತಾರೆ?', level=2)
    doc.add_paragraph('ಭೇಟಿಗಳ ಸಮಯದಲ್ಲಿ, ವಿದ್ಯಾರ್ಥಿಗಳು ಇವುಗಳನ್ನು ಮೌಲ್ಯಮಾಪನ ಮಾಡುತ್ತಾರೆ:', style='List Bullet')
    doc.add_paragraph('ಜನಸಂಖ್ಯಾಶಾಸ್ತ್ರ (Demographics): ಕುಟುಂಬದ ಸಂಯೋಜನೆ, ಎಲ್ಲಾ ಸದಸ್ಯರ ಶಿಕ್ಷಣ, ಉದ್ಯೋಗ.', style='List Bullet')
    doc.add_paragraph('ಸಾಮಾಜಿಕ-ಆರ್ಥಿಕ ಸ್ಥಿತಿ: ಆದಾಯ ಪರಿಶೀಲನೆ (ಆಪ್ ಸ್ವಯಂಚಾಲಿತವಾಗಿ ಬಿ.ಜಿ. ಪ್ರಸಾದ್ ಸ್ಕೇಲ್ ಲೆಕ್ಕಾಚಾರ ಮಾಡುತ್ತದೆ).', style='List Bullet')
    doc.add_paragraph('ಪರಿಸರ ಆರೋಗ್ಯ: ನೀರಿನ ಮೂಲ, ತ್ಯಾಜ್ಯ ವಿಲೇವಾರಿ, ವಾತಾಯನ, ಜನದಟ್ಟಣೆ.', style='List Bullet')
    doc.add_paragraph('ಆರೋಗ್ಯ ವೈಟಲ್ಸ್: ರಕ್ತದೊತ್ತಡ (BP), ನಾಡಿಮಿಡಿತ, BMI, ರಕ್ತದ ಸಕ್ಕರೆ (ಅಗತ್ಯವಿದ್ದರೆ).', style='List Bullet')
    doc.add_paragraph('ತಾಯಿ ಮತ್ತು ಮಕ್ಕಳ ಆರೋಗ್ಯ: ಪ್ರಸವಪೂರ್ವ ಆರೈಕೆ (ANC) ಸ್ಥಿತಿ, ಲಸಿಕೆ ವ್ಯಾಪ್ತಿ.', style='List Bullet')
    
    doc.add_heading('3.2 ವಿದ್ಯಾರ್ಥಿಗಳು ಏನು ಮಾಡುತ್ತಾರೆ (ಚಟುವಟಿಕೆಗಳು)?', level=2)
    doc.add_paragraph('1. ಕುಟುಂಬ ನೋಂದಣಿ: ದತ್ತು ಪಡೆದ ಕುಟುಂಬಗಳಿಗೆ ಡಿಜಿಟಲ್ ಫೋಲ್ಡರ್‌ಗಳನ್ನು ರಚಿಸುವುದು.', style='List Number')
    doc.add_paragraph('2. ನಿಯಮಿತ ಭೇಟಿಗಳು: ಕುಟುಂಬಗಳಿಗೆ ನಿಯತಕಾಲಿಕವಾಗಿ ಭೇಟಿ ನೀಡುವುದು ಮತ್ತು ಆಪ್‌ನಲ್ಲಿ "Family Visits" ಲಾಗ್ ಮಾಡುವುದು.', style='List Number')
    doc.add_paragraph('3. ಆರೋಗ್ಯ ಶಿಕ್ಷಣ: ಮೌಲ್ಯಮಾಪನದ ಆಧಾರದ ಮೇಲೆ ನೈರ್ಮಲ್ಯ, ಪೋಷಣೆ ಮತ್ತು ರೋಗ ತಡೆಗಟ್ಟುವಿಕೆಯ ಬಗ್ಗೆ ಸಲಹೆ ನೀಡುವುದು.', style='List Number')
    doc.add_paragraph('4. ರಿಫ್ಲೆಕ್ಷನ್ ಬರೆಯುವುದು: ಭೇಟಿಯ ನಂತರ, ವಿದ್ಯಾರ್ಥಿಗಳು ಗಿಬ್ಸ್ ಸೈಕಲ್ ಬಳಸಿ ಜರ್ನಲ್ ಬರೆಯುತ್ತಾರೆ. AI ಕೋಚ್ ಇದನ್ನು ವಿಶ್ಲೇಷಿಸಿ ಅವರ ಅನುಭೂತಿ ತಪ್ಪಿ ವೈದ್ಯಕೀಯ ತರ್ಕವನ್ನು ಸುಧಾರಿಸಲು ಸಹಾಯ ಮಾಡುತ್ತದೆ.', style='List Number')

    doc.add_page_break()

    # Section 4: Mentor Workflow
    doc.add_heading('4. ಮೆಂಟರ್ (ಶಿಕ್ಷಕ) ಕೆಲಸದ ಹರಿವು', level=1)
    doc.add_paragraph('ಮೆಂಟರ್‌ಗಳು ವಿದ್ಯಾರ್ಥಿಗಳ ಗುಂಪನ್ನು ಮೇಲ್ವಿಚಾರಣೆ ಮಾಡುತ್ತಾರೆ ಮತ್ತು ಕ್ಷೇತ್ರಕಾರ್ಯದ ಗುಣಮಟ್ಟವನ್ನು ಖಚಿತಪಡಿಸುತ್ತಾರೆ.')
    
    doc.add_heading('4.1 ಮಾನಿಟರಿಂಗ್ ಮತ್ತು ವಿಮರ್ಶೆ', level=2)
    doc.add_paragraph('- ಡ್ಯಾಶ್‌ಬೋರ್ಡ್ ಅವಲೋಕನ: ನಿಯೋಜಿತ ವಿದ್ಯಾರ್ಥಿಗಳ ಪಟ್ಟಿ ಮತ್ತು ಅವರ ಒಟ್ಟು ಕುಟುಂಬಗಳು/ಭೇಟಿಗಳನ್ನು ವೀಕ್ಷಿಸುವುದು.', style='List Bullet')
    doc.add_paragraph('- ರಿಫ್ಲೆಕ್ಷನ್ ಮೌಲ್ಯಮಾಪನ: ವಿದ್ಯಾರ್ಥಿಗಳ ರಿಫ್ಲೆಕ್ಷನ್‌ಗಳನ್ನು ಓದುವುದು ಮತ್ತು ಗ್ರೇಡಿಂಗ್/ಪ್ರತಿಕ್ರಿಯೆ ನೀಡುವುದು. (ಮೆಂಟರ್‌ಗಳಿಗೆ ಸಹಾಯ ಮಾಡಲು ಆಪ್ AI ಒಳನೋಟಗಳನ್ನು ನೀಡುತ್ತದೆ).', style='List Bullet')
    doc.add_paragraph('- ಲಾಗ್ ಪರಿಶೀಲನೆ: ವಿದ್ಯಾರ್ಥಿಗಳು ಲಾಗ್ ಮಾಡಿದ ಭೇಟಿಗಳ ಸತ್ಯಾಸತ್ಯತೆಯನ್ನು ಪರಿಶೀಲಿಸುವುದು.', style='List Bullet')

    doc.add_heading('4.2 ಸಾಮರ್ಥ್ಯ ಮೌಲ್ಯಮಾಪನ (Competency Assessment)', level=2)
    doc.add_paragraph('ವಿದ್ಯಾರ್ಥಿಗಳು ನಿರ್ದಿಷ್ಟ CBME ಸಾಮರ್ಥ್ಯಗಳನ್ನು ಸಾಧಿಸಿದ್ದಾರೆಯೇ ಎಂದು ಮೆಂಟರ್‌ಗಳು ಮೌಲ್ಯಮಾಪನ ಮಾಡುತ್ತಾರೆ (ಉದಾಹರಣೆಗೆ, "ಅನುಭೂತಿಯನ್ನು ಪ್ರದರ್ಶಿಸುವುದು", "ಪೌಷ್ಟಿಕಾಂಶದ ಮೌಲ್ಯಮಾಪನ ನಡೆಸುವುದು").')

    doc.add_page_break()

    # Section 5: Admin Workflow
    doc.add_heading('5. ಅಡ್ಮಿನಿಸ್ಟ್ರೇಟರ್ ಕೆಲಸದ ಹರಿವು', level=1)
    doc.add_paragraph('ಕಾಲೇಜಿನ ಮಟ್ಟದಲ್ಲಿ ಸಿಸ್ಟಮ್ ಸುಗಮವಾಗಿ ನಡೆಯುವುದನ್ನು ಅಡ್ಮಿನ್‌ಗಳು ಖಚಿತಪಡಿಸುತ್ತಾರೆ.')
    
    doc.add_heading('5.1 ಬಳಕೆದಾರ ನಿರ್ವಹಣೆ', level=2)
    doc.add_paragraph('- ಖಾತೆಗಳನ್ನು ರಚಿಸುವುದು: ವಿದ್ಯಾರ್ಥಿಗಳು ಮತ್ತು ಶಿಕ್ಷಕರಿಗಾಗಿ ಖಾತೆಗಳನ್ನು ರಚಿಸುವುದು.', style='List Bullet')
    doc.add_paragraph('- ಮೆಂಟರ್ ನಿಯೋಜನೆ: ವಿದ್ಯಾರ್ಥಿಗಳ ಬ್ಯಾಚ್‌ಗಳನ್ನು ನಿರ್ದಿಷ್ಟ ಮೆಂಟರ್‌ಗಳಿಗೆ ಮ್ಯಾಪ್ ಮಾಡುವುದು.', style='List Bullet')
    
    doc.add_heading('5.2 ಸಿಸ್ಟಮ್ ಮೇಲ್ವಿಚಾರಣೆ', level=2)
    doc.add_paragraph('- ಅನಾಲಿಟಿಕ್ಸ್: ಕಾಲೇಜು ಮಟ್ಟದ ಅಂಕಿಅಂಶಗಳನ್ನು ವೀಕ್ಷಿಸುವುದು (ದತ್ತು ಪಡೆದ ಒಟ್ಟು ಕುಟುಂಬಗಳು, ರೋಗದ ಹೊರೆ ನಕ್ಷೆಗಳು).', style='List Bullet')
    doc.add_paragraph('- ಡೇಟಾ ರಫ್ತು (Export): NMC ವರದಿಗಳು ಅಥವಾ ಸಂಶೋಧನಾ ಉದ್ದೇಶಗಳಿಗಾಗಿ ಡೇಟಾವನ್ನು ರಫ್ತು ಮಾಡುವುದು.', style='List Bullet')

    doc.add_page_break()

    # Section 6: Systematic Feature List
    doc.add_heading('6. ವೈಶಿಷ್ಟ್ಯಗಳ ಪಟ್ಟಿ', level=1)
    
    features = [
        ('1. ದೃಢೀಕರಣ (Authentication)', 'ಸುರಕ್ಷಿತ ಇಮೇಲ್/ಪಾಸ್‌ವರ್ಡ್ ಲಾಗಿನ್, ಪಾತ್ರ-ಆಧಾರಿತ ಪ್ರವೇಶ ನಿಯಂತ್ರಣ (RBAC).'),
        ('2. ಆಫ್‌ಲೈನ್ ಮೋಡ್', 'ಇಂಟರ್ನೆಟ್ ಇಲ್ಲದೆ ಪೂರ್ಣ ಕಾರ್ಯನಿರ್ವಹಣೆ; ಆನ್‌ಲೈನ್‌ಗೆ ಬಂದಾಗ ಆಟೋ-ಸಿಂಕ್.'),
        ('3. ಡಿಜಿಟಲ್ ಫ್ಯಾಮಿಲಿ ಫೋಲ್ಡರ್', 'ಕುಟುಂಬದ ಜನಸಂಖ್ಯಾಶಾಸ್ತ್ರ, SE ಸ್ಥಿತಿ ಮತ್ತು ಆರೋಗ್ಯದ ಸಮಗ್ರ ದಾಖಲೆ.'),
        ('4. ಸ್ವಯಂಚಾಲಿತ ಕ್ಯಾಲ್ಕುಲೇಟರ್‌ಗಳು', 'SE ವರ್ಗೀಕರಣಕ್ಕಾಗಿ ಸಂಯೋಜಿತ ಬಿ.ಜಿ. ಪ್ರಸಾದ್ ಮತ್ತು ಕುಪ್ಪುಸ್ವಾಮಿ ಮಾಪಕಗಳು.'),
        ('5. AI ಮೆಡಿಕಲ್ ಕೋಚ್', 'ವಿದ್ಯಾರ್ಥಿಗಳ ರಿಫ್ಲೆಕ್ಷನ್‌ಗಳನ್ನು ಪರಿಶೀಲಿಸಿ ವೈದ್ಯಕೀಯ ಮಾರ್ಗದರ್ಶನ ನೀಡುವ ಜೆಮಿನಿ AI.'),
        ('6. ಲಾಗ್-ಬುಕ್ ಜನರೇಷನ್', 'ವಿದ್ಯಾರ್ಥಿಗಳಿಗೆ NMC-ಅನುಸರಣೆಯ ಲಾಗ್-ಬುಕ್‌ಗಳ PDF ರಚನೆ.'),
        ('7. ಸಮುದಾಯ ಡ್ಯಾಶ್‌ಬೋರ್ಡ್', 'ಹಳ್ಳಿಯ ಆರೋಗ್ಯ ಸೂಚಕಗಳನ್ನು ತೋರಿಸುವ ಗ್ರಾಫ್‌ಗಳು.'),
        ('8. ಕ್ಲಿನಿಕಲ್ ಮಾರ್ಗಸೂಚಿಗಳು', 'ಭೇಟಿಗಳ ಸಮಯದಲ್ಲಿ ಉಲ್ಲೇಖಕ್ಕಾಗಿ ಪ್ರಮಾಣಿತ ಮಾರ್ಗಸೂಚಿಗಳಿಗೆ (IMNCI, TB, ANC) ಆಫ್‌ಲೈನ್ ಪ್ರವೇಶ.')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ವೈಶಿಷ್ಟ್ಯ'
    hdr_cells[1].text = 'ವಿವರಣೆ'
    
    for feat, desc in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = desc

    # Save
    doc.save('FAP_NextGen_Documentation_Kannada_v2.docx')
    print("Kannada documentation created successfully: FAP_NextGen_Documentation_Kannada_v2.docx")

if __name__ == "__main__":
    try:
        create_kannada_documentation()
    except Exception as e:
        print(f"Error creating documentation: {e}")
