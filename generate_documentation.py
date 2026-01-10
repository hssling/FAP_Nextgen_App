
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_documentation():
    doc = Document()

    # Title Page
    title = doc.add_heading('FAP NextGen App', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Documentation & Role-Based User Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n' * 5)
    
    details = doc.add_paragraph('For: MBBS Students, Faculty Mentors, and Administrators\nContext: Competency-Based Medical Education (CBME)\nFamily Adoption Programme (FAP)')
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Introduction & Objectives')
    doc.add_paragraph('2. Roles & Responsibilities Overview')
    doc.add_paragraph('3. Student Workflow (Assessments & Activities)')
    doc.add_paragraph('4. Mentor/Teacher Workflow')
    doc.add_paragraph('5. Administrator Workflow')
    doc.add_paragraph('6. Systematic Feature List')
    doc.add_paragraph('7. Technical Architecture')
    
    doc.add_page_break()

    # Section 1: Introduction
    doc.add_heading('1. Introduction & Objectives', level=1)
    p = doc.add_paragraph('The ')
    p.add_run('FAP NextGen App').bold = True
    p.add_run(' is a comprehensive digital platform designed to operationalize the Family Adoption Programme (FAP) as per National Medical Commission (NMC) guidelines. It serves as a bridge between medical students, the community, and faculty mentors.')
    
    doc.add_heading('Key Objectives:', level=2)
    doc.add_paragraph('1. To facilitate longitudinal health monitoring of rural families by medical students.', style='List Bullet')
    doc.add_paragraph('2. To provide real-time data for community diagnosis and health interventions.', style='List Bullet')
    doc.add_paragraph('3. To enable mentors to assess student competencies remotely and effectively.', style='List Bullet')

    # Section 2: Roles
    doc.add_heading('2. Roles & Responsibilities Overview', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Primary Responsibility'
    
    roles = [
        ('Student', 'Adopts families, conducts visits, collects data, and writes reflections.'),
        ('Mentor (Teacher)', 'Guides students, reviews reflections, and evaluates performance.'),
        ('Administrator', 'Manages users (students/teachers), oversees system health, creates backups.')
    ]
    
    for role, resp in roles:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = resp

    doc.add_page_break()

    # Section 3: Student Workflow
    doc.add_heading('3. Student Workflow (Assessments & Activities)', level=1)
    doc.add_paragraph('Students are the primary data collectors. Their workflow involves the following systematic steps:')
    
    doc.add_heading('3.1 What Students Assess (Data Collection)', level=2)
    doc.add_paragraph('During visits, students assess:', style='List Bullet')
    doc.add_paragraph('Demographics: Family composition, education, occupation of all members.', style='List Bullet')
    doc.add_paragraph('Socio-Economic Status: Income verification (App auto-calculates BG Prasad Scale).', style='List Bullet')
    doc.add_paragraph('Environmental Health: Water source, waste disposal, ventilation, overcrowding.', style='List Bullet')
    doc.add_paragraph('Health Vitals: Blood Pressure, Pulse, BMI, Blood Sugar (if applicable).', style='List Bullet')
    doc.add_paragraph('Maternal & Child Health: Antenatal care status, Immunization coverage.', style='List Bullet')
    
    doc.add_heading('3.2 What Students Do (Actionable Tasks)', level=2)
    doc.add_paragraph('1. Family Registration: Create digital folders for adopted families.', style='List Number')
    doc.add_paragraph('2. Regular Visits: Visit families periodically (as per schedule) and log "Family Visits" in the app.', style='List Number')
    doc.add_paragraph('3. Health Education: Provide counseling on hygiene, nutrition, and disease prevention based on assessment.', style='List Number')
    doc.add_paragraph('4. Reflection Writing: Post-visit, students write a reflective journal using the Gibbs Cycle. The AI Coach analyzes this to improve their empathy and clinical reasoning.', style='List Number')

    doc.add_page_break()

    # Section 4: Mentor Workflow
    doc.add_heading('4. Mentor (Teacher) Workflow', level=1)
    doc.add_paragraph('Mentors oversee a group of students and ensure quality of fieldwork.')
    
    doc.add_heading('4.1 Monitoring & Review', level=2)
    doc.add_paragraph('- Dashboard Overview: View list of assigned students and their total families/visits.', style='List Bullet')
    doc.add_paragraph('- Reflection Assessment: Read student reflections and provide grading/feedback. (App provides AI insights to help mentors).', style='List Bullet')
    doc.add_paragraph('- Field Log Verification: Verify the authenticity of visits logged by students via geolocation tags (if enabled).', style='List Bullet')

    doc.add_heading('4.2 Competency Assessment', level=2)
    doc.add_paragraph('Mentors evaluate if students have achieved specific CBME competencies (e.g., "Demonstrate empathy", "Conduct nutritional assessment").')

    doc.add_page_break()

    # Section 5: Admin Workflow
    doc.add_heading('5. Administrator Workflow', level=1)
    doc.add_paragraph('Admins are responsible for the smooth running of the respective college\'s instance.')
    
    doc.add_heading('5.1 User Management', level=2)
    doc.add_paragraph('- Create Accounts: Bulk upload or manually create accounts for Students and Teachers.', style='List Bullet')
    doc.add_paragraph('- Assign Mentors: Map batches of students to specific mentors.', style='List Bullet')
    
    doc.add_heading('5.2 System Oversight', level=2)
    doc.add_paragraph('- Analytics: View college-wide statistics (Total families adopted, community disease burden maps).', style='List Bullet')
    doc.add_paragraph('- Data Exports: Export data for NMC reports or research purposes.', style='List Bullet')

    doc.add_page_break()

    # Section 6: Systematic Feature List
    doc.add_heading('6. Systematic Feature List', level=1)
    
    features = [
        ('1. Authentication', 'Secure Email/Password Login, Role-Based Access Control (RBAC).'),
        ('2. Offline Mode', 'Full functionality without internet; auto-sync when online.'),
        ('3. Digital Family Folder', 'Comprehensive record of family demographics, SE status, and health.'),
        ('4. Automated Calculators', 'Integrated BG Prasad & Kuppuswamy scales for SE classification.'),
        ('5. AI Medical Coach', 'Generative AI (Gemini) that reviews student reflections and offers clinical guidance.'),
        ('6. Logbook Generation', 'One-click PDF generation of NMC-compliant logbooks for students.'),
        ('7. Community Dashboard', 'Visual graphs showing village health indicators (e.g., % of hypertension).'),
        ('8. Clinical Guidelines', 'Offline access to standard guidelines (IMNCI, TB, ANC) for reference during visits.')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    
    for feat, desc in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = desc

    doc.add_page_break()
    
    # Section 7: Tech Stack (Brief)
    doc.add_heading('7. Technical Architecture (Brief)', level=1)
    doc.add_paragraph('The app uses a modern tech stack ensuring speed and reliability:')
    doc.add_paragraph('Frontend: React 18 + Vite (PWA)\nBackend: Supabase (PostgreSQL)\nAI: Google Gemini\nHosting: Vercel')

    # Save
    doc.save('FAP_NextGen_Documentation.docx')
    print("Documentation created successfully: FAP_NextGen_Documentation.docx")

if __name__ == "__main__":
    try:
        create_documentation()
    except Exception as e:
        print(f"Error creating documentation: {e}")
