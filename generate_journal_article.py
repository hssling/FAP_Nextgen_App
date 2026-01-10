
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_journal_article():
    doc = Document()

    # Title
    title = doc.add_heading('FAP NextGen: A Digital Innovation for Competency-Based Medical Education under the National Medical Commission Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph('Dr. Siddalingaiah H S')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors.runs[0]
    authors_run.bold = True

    affiliation = doc.add_paragraph('Professor, Department of Community Medicine\nShridevi Institute of Medical Sciences & Research Hospital (SIMS & RH), Tumkur, Karnataka, India')
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = """Background: The introduction of Competency-Based Medical Education (CBME) by the National Medical Commission (NMC) in 2019 necessitated a fundamental shift in how medical students are trained and assessed. A key component of this new curriculum is the Family Adoption Programme (FAP), which mandates longitudinal community-based learning. However, the operationalization of FAP faces significant logistical challenges, including data management, quality assurance of field visits, and providing timely feedback to students. This paper describes the development, features, and potential impact of FAP NextGen, a purpose-built digital application designed to address these challenges.

Methods: FAP NextGen was developed using an iterative, user-centered design process. The application was built on a modern technology stack (React, Supabase, Gemini AI) to ensure scalability, security, and offline functionality critical for rural field settings. The design was informed by NMC guidelines, national health priorities (National Health Policy 2017, Ayushman Bharat), and pedagogical principles of reflective practice (Gibbs' Reflective Cycle).

Results: The resulting application provides a comprehensive digital ecosystem for FAP implementation. Key features include secure role-based access for students, mentors, and administrators; a digital family folder with automated socio-economic classification (BG Prasad Scale); an AI-powered Medical Coach for providing instant feedback on student reflections; and offline-first data synchronization for use in low-connectivity areas. The system generates NMC-compliant digital logbooks, replacing cumbersome paper-based documentation.

Conclusion: FAP NextGen represents a significant step towards modernizing community-based medical education in India. By digitizing the FAP workflow, the application enhances data quality, promotes reflective learning, enables remote mentorship, and aligns with national health digitization efforts. Its adoption can facilitate a more standardized, scalable, and impactful implementation of FAP across Indian medical colleges.

Keywords: CBME, Family Adoption Programme, Medical Education, NMC, Digital Health, AI in Education, Community Medicine."""
    doc.add_paragraph(abstract_text)

    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = """The landscape of medical education in India underwent a paradigm shift with the introduction of Competency-Based Medical Education (CBME) by the National Medical Commission (NMC) in 2019. This outcome-based approach moves away from traditional, knowledge-centric pedagogy towards a model that emphasizes the observable and measurable abilities of a medical graduate.[1] A cornerstone of this reformed curriculum is its renewed focus on community-based learning, recognizing that competent physicians must understand the social determinants of health and the realities of primary care delivery in India.

The Family Adoption Programme (FAP) is the primary vehicle for this community-based learning within the CBME framework. Under FAP, each medical student is required to "adopt" a family from a rural or peri-urban community and follow them longitudinally throughout their undergraduate career. The objectives are multi-fold: to develop empathy, to understand the socio-economic and environmental factors influencing health, to practice clinical skills in a real-world setting, and to contribute to community health improvement through education and basic interventions.

While the intent of FAP is laudable, its implementation presents formidable challenges. The traditional paper-based logbook system is fraught with issues of data quality, traceability, and security. Faculty mentors often lack the time and resources to effectively supervise and provide timely feedback to large numbers of students in the field. The disconnect between the field experience and formal learning can hinder meaningful integration of knowledge. Furthermore, documenting and assessing the "soft" competencies developed through FAP, such as empathy and communication, remains a subjective and often inconsistent process.

This paper presents FAP NextGen, a comprehensive digital application designed to address these implementation challenges. Developed as a public good aligned with the principles of open-source technology, FAP NextGen aims to transform the FAP workflow from a cumbersome administrative task into an enriching, technology-enabled learning experience that is fully aligned with NMC mandates and national health priorities."""

    doc.add_paragraph(intro_text)

    # Alignment with Policies
    doc.add_heading('2. Alignment with NMC, CBME, and National Health Priorities', level=1)

    doc.add_heading('2.1 NMC CBME Curriculum', level=2)
    nmc_text = """The NMC CBME curriculum explicitly mandates the FAP, detailing specific competencies that students must achieve through this exposure. These include competencies related to communication (AETCOM module), community diagnosis, basic epidemiological surveys, and health education. FAP NextGen directly supports the documentation and assessment of these competencies.

The application's "Reflection" module, for instance, is structured around the Gibbs' Reflective Cycle, a pedagogical framework recommended by the NMC for fostering deeper learning from clinical experiences. The integrated AI Medical Coach analyzes these reflections against the CBME competency framework, providing students and mentors with specific feedback on areas such as clinical reasoning, empathy, and professional behavior. This transforms assessment from a purely summative exercise to a continuous, formative process.

Furthermore, the digital logbook generated by FAP NextGen meets the NMC's documentation requirements for submission during internal assessments and university examinations. This standardization ensures comparability across institutions and reduces the subjectivity inherent in manually evaluated paper logbooks.[2]"""
    doc.add_paragraph(nmc_text)

    doc.add_heading('2.2 National Health Policy 2017 and Ayushman Bharat', level=2)
    nhp_text = """The National Health Policy 2017 emphasizes the need to create health systems that are responsive, accessible, and patient-centric, particularly for vulnerable populations.[3] The Family Adoption Programme, when implemented effectively, trains future doctors in exactly this vision—understanding the needs of underserved communities and delivering care within their context.

FAP NextGen aligns with the policy's focus on harnessing digital technology for health. The application contributes to the broader national digital health ecosystem envisioned under the Ayushman Bharat Digital Mission (ABDM). While not directly integrated with ABHA IDs in the current version, the architecture is designed to be interoperable. The longitudinal health data collected on adopted families can, with appropriate consent frameworks, contribute to a community-level health information system that supports evidence-based planning at the Primary Health Centre (PHC) and Community Health Centre (CHC) level.

The app's emphasis on Non-Communicable Disease (NCD) screening aligns with the Ayushman Bharat Health and Wellness Centre initiative, which prioritizes population-level screening for hypertension, diabetes, and common cancers. Students using FAP NextGen are trained to perform these screenings, reinforcing national public health priorities at the grassroots level."""
    doc.add_paragraph(nhp_text)

    doc.add_page_break()

    # Features
    doc.add_heading('3. Key Features of FAP NextGen', level=1)

    doc.add_heading('3.1 Digital Family Folder', level=2)
    folder_text = """The core of FAP NextGen is the Digital Family Folder, a secure, cloud-based repository that replaces the traditional paper file. For each adopted family, students create and maintain a comprehensive record that includes:

•   Demographic Data: Information on all family members, including age, gender, education, and occupation.
•   Socio-Economic Assessment: The application includes an automated calculator for the BG Prasad Scale (updated to 2024), which classifies the family's socio-economic status based on per capita income. This eliminates manual calculation errors and ensures consistent classification.
•   Environmental Health Profile: A structured checklist for documenting housing conditions, water source, sanitation, and waste disposal practices.
•   Longitudinal Health Records: Vitals (blood pressure, weight, blood sugar) and health status of individual members are tracked over time, enabling students to observe trends and the impact of interventions."""
    doc.add_paragraph(folder_text)

    doc.add_heading('3.2 AI-Powered Medical Coach', level=2)
    ai_text = """A distinguishing feature of FAP NextGen is its integration of a Generative AI-powered "Medical Coach." After each field visit, students are prompted to write a reflective journal entry. The AI Coach, powered by Google's Gemini large language model via the OpenRouter API, provides instant, personalized feedback.

The feedback is structured to:
1.  Acknowledge the student's emotional experience and empathy, validating the affective domain of learning.
2.  Identify relevant clinical or public health concepts related to the student's observations (e.g., if a student mentions a coughing patient, the AI might suggest reviewing TB screening protocols).
3.  Offer specific, actionable tips for improvement in areas like communication, clinical examination, or health education techniques.

This AI-assisted feedback loop does not replace the mentor but augments their capacity. Mentors receive AI-analyzed summaries of student reflections, allowing them to focus their limited time on students who need the most guidance and on providing deeper, personalized feedback that only a human educator can offer."""
    doc.add_paragraph(ai_text)

    doc.add_heading('3.3 Offline-First Architecture', level=2)
    offline_text = """A critical design consideration for FAP NextGen was usability in low-resource, low-connectivity environments typical of rural India. The application is built as a Progressive Web App (PWA) with an offline-first architecture.

Students can load the application once while connected to the internet. Subsequently, they can view family records, add new visit notes, and even interact with pre-cached guideline content entirely offline. A proprietary "Base64 Bypass" technology ensures that files (e.g., photos of health records or village infrastructure) are reliably uploaded even on unstable mobile networks by encoding them within the database transaction stream.

All data entered offline is seamlessly synchronized with the central Supabase cloud database as soon as connectivity is restored. This ensures that no data is lost during fieldwork and that the central database always reflects the latest information."""
    doc.add_paragraph(offline_text)

    doc.add_heading('3.4 Role-Based Access and Mentorship', level=2)
    roles_text = """FAP NextGen implements a robust Role-Based Access Control (RBAC) system, recognizing that different stakeholders have different needs and responsibilities:

•   Students: Can create and manage their own family folders, log visits, write reflections, and view clinical guidelines. They cannot see data from other students.
•   Mentors (Teachers): Have a dedicated dashboard displaying all students assigned to them. They can view student family data, read student reflections, provide grades and qualitative feedback, and validate the authenticity of field logs.
•   Administrators: Have full access to manage user accounts (creation, assignment) and view college-wide analytics for reporting to institutional leadership and the NMC.

The Row Level Security (RLS) policies enforced at the database level (Supabase/PostgreSQL) ensure data privacy and integrity, even protecting against potential backend vulnerabilities."""
    doc.add_paragraph(roles_text)

    doc.add_heading('3.5 Integrated Clinical Guidelines', level=2)
    guidelines_text = """To support students during field visits, FAP NextGen includes a built-in repository of standard clinical guidelines. These are accessible offline and cover critical areas relevant to community medicine practice:

•   Antenatal Care (ANC) Guidelines
•   Universal Immunization Programme (UIP) Schedules
•   Integrated Management of Neonatal and Childhood Illnesses (IMNCI)
•   Revised National Tuberculosis Control Programme (RNTCP/NTEP) - TB Screening Algorithms
•   NCD Screening (Hypertension, Diabetes, Cervical Cancer)
•   Family Planning Counseling

This feature ensures that students have evidence-based resources at their fingertips, promoting safe and effective practice even when away from institutional facilities."""
    doc.add_paragraph(guidelines_text)

    doc.add_page_break()

    # Implementation
    doc.add_heading('4. Implementation and Technical Architecture', level=1)
    tech_text = """FAP NextGen is built on a modern, scalable technology stack designed for reliability and security:

•   Frontend: React 18 with Vite, providing a fast, responsive single-page application experience. The UI is optimized for mobile devices, the primary means of access for students in the field.
•   Backend: Supabase, an open-source Backend-as-a-Service (BaaS) platform, provides a PostgreSQL database, authentication, and secure file storage. Supabase's built-in Row Level Security (RLS) is critical for enforcing data access policies.
•   AI Integration: The AI Medical Coach utilizes Google's Gemini model, accessed via the OpenRouter API, for analyzing student reflections and generating feedback.
•   Hosting: The application is deployed on Vercel, a global edge network, ensuring low latency access from anywhere in India.

The application is open-source and designed for easy deployment by individual medical institutions. A comprehensive documentation package, including database schema scripts and deployment guides for Vercel and Supabase, is provided."""
    doc.add_paragraph(tech_text)

    # Discussion
    doc.add_heading('5. Discussion and Potential Impact', level=1)
    discussion_text = """The digitization of the Family Adoption Programme through FAP NextGen addresses several long-standing challenges in community-based medical education.

First, it significantly improves data quality and integrity. The structured forms, automated calculations, and digital audit trails eliminate the ambiguities and potential for fabrication inherent in handwritten logbooks. This enhances the reliability of student work and the value of the data collected.

Second, the AI Medical Coach represents a novel application of generative AI in medical education. By providing instant, personalized feedback, it addresses the critical bottleneck of faculty time. In India, where student-to-faculty ratios are often high, such technological augmentation can democratize access to quality mentorship. It is important to note that this AI is designed as an augmentation tool, not a replacement for the irreplaceable human connection between mentor and student.

Third, the offline-first architecture ensures equity. Students posted to remote villages are not disadvantaged compared to those in peri-urban areas with better connectivity. This is crucial for ensuring that the educational experience of FAP is uniform across diverse geographical settings.

Fourth, the application contributes to the broader national goal of creating a digitally literate health workforce. By using FAP NextGen, students become familiar with concepts like electronic health records, data privacy, and digital consent—skills essential for practicing medicine in an increasingly digital future.

Finally, the longitudinal data collected through FAP NextGen, when aggregated and anonymized, can provide valuable insights into community health trends. This data can inform public health research and health system planning at local and regional levels, creating a virtuous cycle between medical education and public health improvement."""
    doc.add_paragraph(discussion_text)

    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    conclusion_text = """FAP NextGen is a purpose-built digital solution designed to operationalize the Family Adoption Programme mandated by the National Medical Commission's CBME curriculum. By addressing the practical challenges of paper-based documentation, remote mentorship, and offline access, the application has the potential to significantly enhance the quality and impact of community-based medical education in India.

The integration of AI for reflective practice feedback represents an innovative frontier in medical pedagogy, promising to scale personalized learning while respecting the indispensable role of human mentors. As India continues its journey towards health system strengthening and digital transformation, tools like FAP NextGen can play a pivotal role in preparing a new generation of physicians who are technologically adept, community-oriented, and competent to meet the nation's health challenges.

Future work will focus on multi-institutional pilot studies to gather user feedback, explore integration with the ABDM ecosystem, and develop more sophisticated analytics dashboards for institutional decision-making."""
    doc.add_paragraph(conclusion_text)

    # References
    doc.add_heading('References', level=1)
    references = [
        '1. National Medical Commission. Graduate Medical Education Regulations (Competency Based Medical Education). New Delhi: Gazette of India; 2019.',
        '2. Supe A, Burdick W. Competency-Based Medical Education in Indian Undergraduate Medical Education. Indian J Community Med. 2020;45(Suppl 1):S15-S18.',
        '3. Ministry of Health and Family Welfare, Government of India. National Health Policy 2017. New Delhi: MoHFW; 2017.',
        '4. Gibbs G. Learning by Doing: A Guide to Teaching and Learning Methods. Oxford: Further Education Unit, Oxford Polytechnic; 1988.',
        '5. NITI Aayog. Ayushman Bharat Digital Mission: Strategy Document. New Delhi: Government of India; 2021.'
    ]
    for ref in references:
        doc.add_paragraph(ref)

    # Save
    doc.save('FAP_NextGen_Journal_Article.docx')
    print("Journal article created successfully: FAP_NextGen_Journal_Article.docx")

if __name__ == "__main__":
    try:
        create_journal_article()
    except Exception as e:
        print(f"Error creating article: {e}")
