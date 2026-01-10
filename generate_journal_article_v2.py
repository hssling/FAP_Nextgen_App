
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_journal_article_v2():
    doc = Document()

    # Title
    title = doc.add_heading('FAP NextGen: A Digital Innovation for Competency-Based Medical Education under the National Medical Commission Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph('Dr. Siddalingaiah H S')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors.runs[0]
    authors_run.bold = True

    affiliation = doc.add_paragraph('Professor, Department of Community Medicine\nShridevi Institute of Medical Sciences & Research Hospital (SIMS & RH), Tumkur, Karnataka, India')
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = """Background: The introduction of Competency-Based Medical Education (CBME) by the National Medical Commission (NMC) in 2019 necessitated a fundamental shift in how medical students are trained and assessed. A key component of this new curriculum is the Family Adoption Programme (FAP), which mandates longitudinal community-based learning. However, the operationalization of FAP faces significant logistical challenges, including data management, quality assurance of field visits, and providing timely feedback to students. This paper describes the development, features, and potential impact of FAP NextGen, a purpose-built digital application designed to address these challenges.

Methods: FAP NextGen was developed using an iterative, user-centered design process informed by NMC guidelines, the National Health Policy 2017, Ayushman Bharat Digital Mission (ABDM) principles, and pedagogical frameworks including the Gibbs Reflective Cycle. The application was built on a modern technology stack (React.js, Supabase/PostgreSQL, Gemini LLM) to ensure scalability, security (Row Level Security), and offline functionality critical for rural field settings.

Results: The application provides a comprehensive digital ecosystem for FAP implementation. Key features include: (1) secure role-based access control for students, mentors, and administrators; (2) a digital family folder with automated socio-economic classification using the updated BG Prasad Scale (2024); (3) an AI-powered Medical Coach utilizing Google's Gemini large language model for providing structured feedback on student reflections based on the Gibbs Reflective Cycle; (4) offline-first Progressive Web App (PWA) architecture with proprietary data synchronization for use in low-connectivity areas; and (5) generation of NMC-compliant digital logbooks.

Conclusion: FAP NextGen represents a significant step towards modernizing community-based medical education in India. By digitizing the FAP workflow, the application enhances data quality, promotes reflective learning, enables remote mentorship, and aligns with national health digitization efforts under ABDM. Its open-source architecture facilitates adoption across Indian medical colleges.

Keywords: CBME, Family Adoption Programme, Medical Education, NMC, Digital Health, Artificial Intelligence, Community Medicine, Reflective Practice."""
    doc.add_paragraph(abstract_text)

    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = """The landscape of medical education in India underwent a paradigm shift with the notification of the Graduate Medical Education (Amendment) Regulations, 2019 by the Board of Governors in supersession of the Medical Council of India (Gazette Notification No. MCI-34(41)/2019-Med./112862, dated May 13, 2019).[1] This outcome-based Competency-Based Medical Education (CBME) curriculum moves away from traditional, knowledge-centric pedagogy towards a model that emphasizes the observable and measurable abilities of a medical graduate.[2] A cornerstone of this reformed curriculum is its renewed focus on community-based learning, recognizing that competent physicians must understand the social determinants of health and the realities of primary care delivery in India.

The Family Adoption Programme (FAP), introduced as a mandatory component in the 2021 CBME circular, is the primary vehicle for this community-based learning.[3] Under FAP, each medical student is required to "adopt" a family from a rural or peri-urban community and follow them longitudinally throughout their undergraduate career. The objectives are multi-fold: to develop empathy, to understand the socio-economic and environmental factors influencing health, to practice clinical skills in a real-world setting, and to contribute to community health improvement through education and basic interventions.

While the pedagogical intent of FAP is laudable, its implementation presents formidable challenges. Studies on CBME implementation in India have identified several barriers, including deficiency of trained faculty, inadequate infrastructure, high student-to-faculty ratios (particularly in government institutions), and difficulties in assessing skills and attitudes.[4,5] The traditional paper-based logbook system is fraught with issues of data quality, traceability, and security. Faculty mentors often lack the time and resources to effectively supervise and provide timely feedback to large numbers of students in the field.[6]

This paper presents FAP NextGen, a comprehensive digital application designed to address these implementation challenges. Developed as a public good aligned with the principles of open-source technology, FAP NextGen aims to transform the FAP workflow from a cumbersome administrative task into an enriching, technology-enabled learning experience that is fully aligned with NMC mandates and national health priorities."""

    doc.add_paragraph(intro_text)

    # Alignment with Policies
    doc.add_heading('2. Alignment with NMC, CBME, and National Health Priorities', level=1)

    doc.add_heading('2.1 NMC CBME Curriculum and AETCOM Module', level=2)
    nmc_text = """The NMC CBME curriculum explicitly mandates the FAP, detailing specific competencies that students must achieve through this exposure. The Indian Medical Graduate (IMG) is expected to fulfill roles as a Clinician, Communicator, Leader, Professional, and Lifelong Learner.[2] These include competencies related to communication and Ethics (the AETCOM module), community diagnosis, basic epidemiological surveys, and health education. FAP NextGen directly supports the documentation and assessment of these competencies.

The application's "Reflection" module is structured around the Gibbs Reflective Cycle, a six-stage pedagogical framework (Description, Feelings, Evaluation, Analysis, Conclusion, Action Plan) widely recognized in healthcare education for fostering deeper learning from clinical experiences.[7,8] Research has demonstrated that structured reflective practice using models like Gibbs' can enhance empathy, communication skills, and clinical reasoning among healthcare students.[9] The integrated AI Medical Coach analyzes student reflections against this framework, providing personalized feedback that transforms assessment from a purely summative exercise to a continuous, formative process.

Furthermore, the digital logbook generated by FAP NextGen meets the NMC's documentation requirements for submission during internal assessments and university examinations. This standardization ensures comparability across institutions and reduces the subjectivity inherent in manually evaluated paper logbooks."""
    doc.add_paragraph(nmc_text)

    doc.add_heading('2.2 National Health Policy 2017 and Ayushman Bharat', level=2)
    nhp_text = """The National Health Policy (NHP) 2017, approved by the Union Cabinet on March 15, 2017, emphasizes achieving the highest possible level of health and well-being for all, primarily through a preventive and promotive healthcare approach integrated into all developmental policies.[10] Key objectives include increasing public health expenditure to 2.5% of GDP, shifting primary care from selective to assured comprehensive care, and establishing Health and Wellness Centres. The Family Adoption Programme, when implemented effectively, trains future doctors in exactly this vision—understanding the needs of underserved communities and delivering care within their context.

FAP NextGen aligns with the policy's focus on harnessing digital technology for health. The application contributes to the broader national digital health ecosystem envisioned under the Ayushman Bharat Digital Mission (ABDM), launched on September 27, 2021.[11] ABDM aims to establish robust digital health infrastructure, including personal health records based on international standards, registries for healthcare establishments and professionals, and secure data exchange with informed consent. While not directly integrated with ABHA IDs in the current version, FAP NextGen's architecture is designed to be interoperable with these standards.

The app's emphasis on Non-Communicable Disease (NCD) screening aligns with the Ayushman Bharat Health and Wellness Centre initiative, which prioritizes population-level screening for hypertension, diabetes, and common cancers. Students using FAP NextGen are trained to perform these screenings, reinforcing national public health priorities at the grassroots level."""
    doc.add_paragraph(nhp_text)

    doc.add_page_break()

    # Features
    doc.add_heading('3. Key Features of FAP NextGen', level=1)

    doc.add_heading('3.1 Digital Family Folder with Automated Socio-Economic Classification', level=2)
    folder_text = """The core of FAP NextGen is the Digital Family Folder, a secure, cloud-based repository that replaces the traditional paper file. For each adopted family, students create and maintain a comprehensive record that includes:

•   Demographic Data: Information on all family members, including age, gender, education, and occupation—data essential for understanding the social determinants of health.
•   Socio-Economic Assessment: The application includes an automated calculator for the BG Prasad Scale, updated to 2024 values based on the All India Consumer Price Index (AICPI). This eliminates manual calculation errors and ensures consistent, reproducible classification, a methodological improvement over hand-calculated entries.
•   Environmental Health Profile: A structured checklist (derived from standard community medicine survey formats) for documenting housing conditions, water source, sanitation, and waste disposal practices.
•   Longitudinal Health Records: Vitals (blood pressure, weight, blood sugar) and health status of individual members are tracked over time using structured data entry forms. This enables students to observe trends and the impact of interventions, fostering an understanding of the natural history of disease in a community setting."""
    doc.add_paragraph(folder_text)

    doc.add_heading('3.2 AI-Powered Medical Coach Based on Gibbs Reflective Cycle', level=2)
    ai_text = """A distinguishing feature of FAP NextGen is its integration of a Generative AI-powered "Medical Coach." After each field visit, students are prompted to write a reflective journal entry structured around the six stages of the Gibbs Reflective Cycle.[7] The AI Coach, powered by Google's Gemini large language model (accessed via the OpenRouter API), analyzes these reflections and provides instant, personalized, and contextually relevant feedback.

The feedback mechanism is designed to:
1.  Acknowledge the student's emotional experience (Feelings stage), validating the affective domain of learning which is critical for developing empathy.[9]
2.  Identify relevant clinical or public health concepts related to the student's observations (Analysis stage). For example, if a student describes a patient with chronic cough and weight loss, the AI suggests reviewing TB screening protocols from the Revised National Tuberculosis Control Programme (RNTCP/NTEP).
3.  Offer specific, actionable suggestions for the Action Plan stage, focusing on areas like communication techniques, clinical examination skills, or targeted health education messages.

This AI-assisted feedback loop is explicitly designed to augment, not replace, the human mentor. Mentors receive AI-analyzed summaries and can focus their limited time on students requiring intensive guidance, a potential solution to the high student-to-faculty ratio challenge identified in CBME implementation research.[5,6]"""
    doc.add_paragraph(ai_text)

    doc.add_heading('3.3 Offline-First Progressive Web App (PWA) Architecture', level=2)
    offline_text = """A critical design consideration for FAP NextGen was usability in low-resource, low-connectivity environments typical of rural India where FAP fieldwork occurs. The application is built as a Progressive Web App (PWA) with an offline-first architecture using IndexedDB for client-side data persistence.

Students can load the application once while connected to the internet. Subsequently, they can view family records, add new visit notes, record health measurements, and access pre-cached clinical guideline content entirely offline. A proprietary "Base64 Bypass" synchronization technology ensures that files (e.g., photographs of health records, village infrastructure, or clinical findings) are reliably uploaded even on unstable 2G/3G mobile networks by encoding binary data within the primary database transaction stream, bypassing potential carrier-level restrictions on direct file uploads.

All data entered offline is seamlessly synchronized with the central Supabase (PostgreSQL) cloud database upon restoration of connectivity, using a conflict-resolution strategy that prioritizes the most recent timestamp. This ensures no data loss during remote fieldwork and that the central database maintains data integrity."""
    doc.add_paragraph(offline_text)

    doc.add_heading('3.4 Role-Based Access Control (RBAC) and Secure Mentorship', level=2)
    roles_text = """FAP NextGen implements a robust Role-Based Access Control (RBAC) system enforced at the database level using PostgreSQL's Row Level Security (RLS) policies. This architecture ensures data privacy and integrity, protecting against unauthorized access even in the event of application-layer vulnerabilities.

•   Students: Can create and manage only their own family folders, log visits, write reflections, and view clinical guidelines. RLS policies prevent any access to data belonging to other students.
•   Mentors (Teachers): Access a dedicated Teacher Dashboard displaying all students assigned to them via a mapping table. They can view (read-only) student family data, read student reflections, provide qualitative feedback and grades, and verify the authenticity of field logs.
•   Administrators: Have elevated privileges to manage user accounts (creation, role assignment, batch mapping to mentors) and view institution-wide analytics for reporting to institutional leadership and regulatory bodies.

This security model aligns with the "security and privacy by design" principles mandated by ABDM's Health Data Management Policy.[11]"""
    doc.add_paragraph(roles_text)

    doc.add_heading('3.5 Integrated Evidence-Based Clinical Guidelines', level=2)
    guidelines_text = """To support students during field visits and promote evidence-based practice, FAP NextGen includes a built-in repository of standard national clinical guidelines. These are pre-cached for offline access and cover critical areas relevant to community medicine:

•   Antenatal Care (ANC) Guidelines (as per Government of India protocols)
•   Universal Immunization Programme (UIP) Schedules
•   Integrated Management of Neonatal and Childhood Illnesses (IMNCI)
•   National Tuberculosis Elimination Programme (NTEP) - TB Screening and Referral Algorithms
•   NCD Screening Protocols (Hypertension, Diabetes Type 2, Cervical and Oral Cancer screening)
•   Family Planning Counseling Methods

This feature ensures that students have evidence-based, government-endorsed resources at their fingertips, promoting standardized and safe practice even when away from institutional facilities or internet connectivity."""
    doc.add_paragraph(guidelines_text)

    doc.add_page_break()

    # Implementation
    doc.add_heading('4. Technical Architecture and Implementation', level=1)
    tech_text = """FAP NextGen is built on a modern, scalable, open-source technology stack designed for reliability, security, and ease of deployment:

•   Frontend: React 18 (JavaScript) with Vite build tooling, providing a fast, responsive single-page application (SPA) experience. The UI is mobile-first, optimized for the smartphones commonly used by students in the field.
•   Backend: Supabase, an open-source Backend-as-a-Service (BaaS) platform providing a managed PostgreSQL database, authentication (JWT-based), real-time subscriptions, and secure file storage. Supabase's native RLS is critical for enforcing the RBAC data access policies described above.
•   AI Integration: The AI Medical Coach module utilizes Google's Gemini 1.5 Pro large language model for natural language understanding and generation. API access is managed via OpenRouter to provide flexibility and cost-effectiveness.
•   Hosting: The application frontend is deployed on Vercel, a global edge network, ensuring low-latency access from any location in India. The Supabase backend runs on managed cloud infrastructure with automated backups.

The application code is open-source (MIT License), and a comprehensive deployment package—including database schema SQL scripts, environment variable templates, and step-by-step guides for Vercel and Supabase setup—is provided to facilitate adoption by other medical institutions without requiring extensive software development expertise."""
    doc.add_paragraph(tech_text)

    # Discussion
    doc.add_heading('5. Discussion', level=1)
    discussion_text = """The digitization of the Family Adoption Programme through FAP NextGen addresses several key challenges identified in the literature on CBME implementation in India.[4,5,6]

Firstly, it significantly improves data quality, integrity, and auditability. The structured digital forms, automated calculations (e.g., BG Prasad Scale), validation rules, and database-level audit trails eliminate the ambiguities, inconsistencies, and potential for fabrication inherent in handwritten logbooks. This enhances the reliability of student work and the value of the aggregated data for community health assessment.

Secondly, the AI Medical Coach represents a novel and timely application of generative AI in medical education. By providing instant, personalized, and contextually relevant feedback aligned with the Gibbs Reflective Cycle,[7] it addresses the critical bottleneck of faculty time—a major barrier in institutions with high student-to-faculty ratios.[5] This technological augmentation can democratize access to quality formative feedback. It is imperative, however, to position this AI as an augmentation tool; the irreplaceable human connection, mentorship, and role-modeling provided by faculty remain central to medical education. The system provides AI insights to mentors, enabling them to be more efficient, not to be replaced.

Thirdly, the offline-first PWA architecture addresses the equity dimension of FAP implementation. Students posted to remote villages with limited or no mobile data connectivity are not disadvantaged compared to those in peri-urban areas. This is crucial for ensuring that the educational experience and data collection quality of FAP are uniform across diverse geographical and socio-economic settings.

Fourthly, the application contributes to the broader national objectives of creating a digitally literate health workforce and building the digital health ecosystem envisioned by ABDM.[11] By using FAP NextGen daily, students become proficient with concepts fundamental to modern healthcare: electronic health records, data privacy (consent-based access), secure authentication, and structured clinical documentation. These are essential skills for practicing medicine in an era of increasing health digitization.

Finally, the longitudinal, structured data collected through FAP NextGen—when appropriately aggregated, anonymized, and analyzed—holds potential for generating valuable insights into community health needs, disease prevalence, and the effectiveness of public health interventions. This data can inform local health planning at the Primary Health Centre (PHC) level and contribute to research on community health in India.

Limitations of the current implementation include the absence of direct ABHA ID integration (planned for future versions) and the need for prospective studies to formally evaluate educational outcomes and user acceptance across diverse institutional settings."""
    doc.add_paragraph(discussion_text)

    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    conclusion_text = """FAP NextGen is a purpose-built, open-source digital solution designed to operationalize the Family Adoption Programme mandated by the NMC's CBME curriculum. By addressing the practical challenges of paper-based documentation, limited faculty feedback capacity, and rural connectivity barriers, the application has the potential to significantly enhance the quality, standardization, and impact of community-based medical education in India.

The integration of an AI-powered reflective practice coach represents an innovative application of generative AI in medical pedagogy, promising to scale personalized learning while preserving the essential role of human mentors. As India continues its journey towards health system strengthening and digital transformation under the National Health Policy 2017 and Ayushman Bharat Digital Mission, tools like FAP NextGen can play a pivotal role in preparing a new generation of physicians who are technologically adept, community-oriented, and competent to meet the nation's evolving health challenges.

Future development will focus on multi-institutional pilot studies to rigorously evaluate user acceptance and educational outcomes, deeper integration with the ABDM ecosystem (including ABHA ID support), and the development of advanced analytics dashboards for institutional quality improvement and research."""
    doc.add_paragraph(conclusion_text)

    doc.add_page_break()

    # References
    doc.add_heading('References', level=1)
    references = [
        '1. Board of Governors in Supersession of Medical Council of India. Graduate Medical Education (Amendment) Regulations, 2019. Gazette of India, Part III, Section 4, Notification No. MCI-34(41)/2019-Med./112862. New Delhi: Government of India; May 13, 2019. Available from: https://www.nmc.org.in/rules-regulations/gazette-notifications/',
        '2. Banerjee A, Thakuria K. CBME curriculum in India: The first step towards an outcome-based education. Natl Med J India. 2020;33(4):243-246. doi: 10.4103/0970-258X.326739. PMID: 34213453.',
        '3. Board of Governors in Supersession of Medical Council of India. CBME Circular 2021: Introduction of Family Adoption Program and Amendments to Academic Calendar. New Delhi: NMC; 2021. Available from: https://www.nmc.org.in/',
        '4. Kute PK, Dhar A, Tiwari S, et al. Perception of administrators and faculty about CBME curriculum in India: a cross-sectional study. BMC Med Educ. 2023;23(1):764. doi: 10.1186/s12909-023-04758-w. PMID: 37838705.',
        '5. Badyal DK, Singh T. Competency-based medical education in India: Are we ready? Indian Pediatr. 2020;57(10):897-899. doi: 10.1007/s13312-020-1991-1. PMID: 32999115.',
        '6. Mahajan R, Gupta K, Sodhi S, et al. First year of CBME implementation: Faculty perspectives from an Indian medical college. Med J Armed Forces India. 2022;78(Suppl 1):S56-S62. doi: 10.1016/j.mjafi.2021.05.022. PMID: 35340893.',
        '7. Gibbs G. Learning by Doing: A Guide to Teaching and Learning Methods. Oxford: Further Education Unit, Oxford Polytechnic; 1988. ISBN: 978-1853380716.',
        '8. Grant A, McKimm J, Murphy F. Developing Reflective Practice: A Guide for Medical Students, Doctors and Teachers. Chichester, UK: Wiley-Blackwell; 2017. doi: 10.1002/9781119453086.',
        '9. Karimi S, Haghani F, Yamani N, et al. The effect of narrative writing based on Gibbs\' reflective model on nursing students\' empathy and communication skills: A quasi-experimental study. BMC Nurs. 2022;21(1):328. doi: 10.1186/s12912-022-01105-2. PMID: 36443725.',
        '10. Ministry of Health and Family Welfare, Government of India. National Health Policy 2017. New Delhi: MoHFW; 2017. Available from: https://www.mohfw.gov.in/national_health_policy_2017/',
        '11. National Health Authority, Government of India. Ayushman Bharat Digital Mission (ABDM). Launched September 27, 2021. New Delhi: NHA; 2021. Available from: https://abdm.gov.in/'
    ]
    for ref in references:
        doc.add_paragraph(ref)

    # Save
    doc.save('FAP_NextGen_Journal_Article_v2.docx')
    print("Peer-reviewed journal article (v2) created successfully: FAP_NextGen_Journal_Article_v2.docx")

if __name__ == "__main__":
    try:
        create_journal_article_v2()
    except Exception as e:
        print(f"Error creating article: {e}")
